from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/Users/Lenovo/Downloads/Project_Report_Template.docx")
OUTPUT = ROOT / "report" / "ScholarSynth_AI_Final_Report.docx"
ASSETS = ROOT / "report" / "assets_final"

TITLE = (
    "Autonomous Research Assistant: A Multi-Agent Generative AI System for "
    "Research Paper Exploration, Literature Review Generation, and Research Gap Analysis"
)

TEAM = [
    ("Akshit Wadhwa", "230784"),
    ("Anish Choudhary", "230820"),
    ("Deepnder", "230821"),
]

ACCENT = RGBColor(22, 78, 121)
MUTED = RGBColor(74, 92, 112)
GREEN = RGBColor(26, 127, 94)
GOLD = RGBColor(151, 103, 13)
BLUE_HEX = "164E79"
LIGHT_BLUE = "EAF3FB"
LIGHT_GREEN = "E8F7F1"
LIGHT_GOLD = "FFF3D7"
LIGHT_RED = "FDECEF"
TABLE_COUNTER = 0


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    names = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for name in names:
        path = Path(name)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = str(text).split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def clear_document_body(document: Document) -> None:
    body = document._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_document(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = ACCENT
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.color.rgb = ACCENT


def add_body(document: Document, text: str, color: RGBColor | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    if color:
        r.font.color.rgb = color


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run("- ").bold = True
        p.add_run(item).font.size = Pt(10.2)


def add_numbered(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(f"{index}. ").bold = True
        p.add_run(item).font.size = Pt(10.2)


def add_callout(document: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.5)
    br = p.add_run(body)
    br.font.size = Pt(10)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    global TABLE_COUNTER
    TABLE_COUNTER += 1
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = ASSETS / f"table_{TABLE_COUNTER:02d}.png"

    image_width = 1600
    margin_x = 45
    pad_x = 22
    pad_y = 14
    header_font = font(26, True)
    body_font = font(23)
    line_gap = 7

    if widths:
        total = sum(widths)
        col_widths = [int((image_width - margin_x * 2) * width / total) for width in widths]
    else:
        col_widths = [int((image_width - margin_x * 2) / len(headers))] * len(headers)
    col_widths[-1] += image_width - margin_x * 2 - sum(col_widths)

    scratch = Image.new("RGB", (image_width, 200), "white")
    draw = ImageDraw.Draw(scratch)

    def row_height(values: list[str], fnt: ImageFont.FreeTypeFont) -> int:
        max_lines = 1
        for value, col_width in zip(values, col_widths):
            max_lines = max(max_lines, len(wrap(draw, str(value), fnt, col_width - 2 * pad_x)))
        line_height = fnt.getbbox("Ag")[3] - fnt.getbbox("Ag")[1] + line_gap
        return max(58, pad_y * 2 + max_lines * line_height)

    all_rows = [headers] + rows
    heights = [row_height(headers, header_font)] + [row_height(row, body_font) for row in rows]
    image_height = margin_x * 2 + sum(heights)
    img = Image.new("RGB", (image_width, image_height), "white")
    draw = ImageDraw.Draw(img)

    y = margin_x
    for row_index, (row, height) in enumerate(zip(all_rows, heights)):
        x = margin_x
        is_header = row_index == 0
        fill = "#164E79" if is_header else ("#F7FBFF" if row_index % 2 == 0 else "#FFFFFF")
        text_fill = "#FFFFFF" if is_header else "#1F2A36"
        fnt = header_font if is_header else body_font
        for value, col_width in zip(row, col_widths):
            draw.rectangle((x, y, x + col_width, y + height), fill=fill, outline="#B7C7D4", width=2)
            lines = wrap(draw, str(value), fnt, col_width - 2 * pad_x)
            line_height = fnt.getbbox("Ag")[3] - fnt.getbbox("Ag")[1] + line_gap
            text_y = y + max((height - len(lines) * line_height) // 2, pad_y)
            for line in lines:
                draw.text((x + pad_x, text_y), line, fill=text_fill, font=fnt)
                text_y += line_height
            x += col_width
        y += height
    draw.rectangle((margin_x, margin_x, image_width - margin_x, image_height - margin_x), outline="#8AA6BF", width=3)
    img.save(out)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(out), width=Inches(6.75))


def read_csv_dict(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def line_count(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open("r", encoding="utf-8", errors="ignore") as fh:
        return sum(1 for _ in fh)


def project_numbers() -> dict[str, str]:
    return {
        "raw_papers": f"{max(line_count(ROOT / 'data/raw_papers.csv') - 1, 0):,}",
        "processed_chunks": f"{max(line_count(ROOT / 'data/processed_papers.csv') - 1, 0):,}",
        "finetune_examples": f"{line_count(ROOT / 'data/finetune_dataset.jsonl'):,}",
        "train_chunks": f"{max(line_count(ROOT / 'data/train.csv') - 1, 0):,}",
        "val_chunks": f"{max(line_count(ROOT / 'data/val.csv') - 1, 0):,}",
        "test_chunks": f"{max(line_count(ROOT / 'data/test.csv') - 1, 0):,}",
        "train_ft": f"{line_count(ROOT / 'data/finetune_train.jsonl'):,}",
        "val_ft": f"{line_count(ROOT / 'data/finetune_val.jsonl'):,}",
        "test_ft": f"{line_count(ROOT / 'data/finetune_test.jsonl'):,}",
    }


def create_workflow_image() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = ASSETS / "scholarsynth_workflow.png"
    w, h = 1800, 980
    img = Image.new("RGB", (w, h), "#F7FAFC")
    draw = ImageDraw.Draw(img)
    title_f = font(48, True)
    sub_f = font(24)
    head_f = font(24, True)
    body_f = font(18)

    draw.rounded_rectangle((35, 32, w - 35, 158), radius=24, fill="#FFFFFF", outline="#D7E2EA", width=2)
    draw.text((70, 55), "ScholarSynth AI: End-to-End System Workflow", fill="#123B5D", font=title_f)
    draw.text((72, 112), "Paper retrieval + semantic storage + RAG + LoRA fine-tuning + guardrails + Streamlit UI", fill="#4A6178", font=sub_f)

    boxes = [
        (70, 230, 345, 360, "#DFF3FF", "User Topic", "Research question from Streamlit UI"),
        (420, 230, 735, 360, "#EAE7FF", "Paper Search Agent", "arXiv and Semantic Scholar APIs"),
        (810, 230, 1115, 360, "#E9F8EF", "Preprocessing", "clean, normalize, deduplicate, chunk"),
        (1190, 230, 1520, 360, "#FFF3D7", "Application Dataset", "5,000 papers; 10,339 chunks"),
        (165, 500, 480, 635, "#E9F8EF", "Storage Layer", "SQLite metadata + ChromaDB vectors"),
        (560, 500, 875, 635, "#DFF3FF", "Retriever", "all-MiniLM-L6-v2 semantic search"),
        (955, 500, 1270, 635, "#EAE7FF", "Generator", "FLAN-T5-base, prompts, RAG context"),
        (1350, 500, 1665, 635, "#FDECEF", "PEFT LoRA", "fine-tuned academic adapter"),
        (370, 780, 715, 900, "#F2ECFF", "Multi-Agent Tasks", "literature review, Q&A, gaps, explanations"),
        (795, 780, 1140, 900, "#E8F7F1", "Guardrails", "input validation and output grounding checks"),
        (1220, 780, 1585, 900, "#FFF3D7", "Evaluation", "BLEU, ROUGE, BERTScore-ready, error analysis"),
    ]

    for x1, y1, x2, y2, fill, heading, body in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline="#AFC0CD", width=3)
        draw.text((x1 + 22, y1 + 22), heading, fill="#123B5D", font=head_f)
        y = y1 + 65
        for line in wrap(draw, body, body_f, x2 - x1 - 44):
            draw.text((x1 + 22, y), line, fill="#36546D", font=body_f)
            y += 25

    arrows = [
        ((350, 295), (415, 295)), ((740, 295), (805, 295)), ((1120, 295), (1185, 295)),
        ((1355, 365), (1510, 495)), ((1190, 565), (880, 565)), ((555, 565), (485, 565)),
        ((880, 565), (950, 565)), ((1275, 565), (1345, 565)), ((1510, 640), (1510, 775)),
        ((1110, 640), (975, 775)), ((710, 640), (550, 775)), ((720, 840), (790, 840)),
        ((1145, 840), (1215, 840)),
    ]
    for (sx, sy), (ex, ey) in arrows:
        draw.line((sx, sy, ex, ey), fill="#2D73A9", width=6)
        dx = 1 if ex >= sx else -1
        draw.polygon([(ex, ey), (ex - 20 * dx, ey - 11), (ex - 20 * dx, ey + 11)], fill="#2D73A9")

    img.save(out)
    return out


def create_metrics_chart() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = ASSETS / "model_comparison_chart.png"
    rows = read_csv_dict(ROOT / "outputs/baseline_200_comparison_table.csv")
    rows = sorted(rows, key=lambda r: float(r.get("rougeL", 0)), reverse=True)
    w, h = 1500, 820
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(24, True)
    small_f = font(22)
    draw.text((60, 40), "Final Model Comparison on 200 Evaluation Examples", fill="#123B5D", font=title_f)
    draw.text((62, 94), "Fine-tuned LoRA gives the strongest lexical match; RAG + LoRA adds source grounding.", fill="#4A6178", font=small_f)

    metrics = [("sacrebleu", "BLEU", "#2D73A9"), ("rouge1", "ROUGE-1", "#28A06A"), ("rougeL", "ROUGE-L", "#C18616")]
    max_values = {"sacrebleu": max(float(r["sacrebleu"]) for r in rows) or 1, "rouge1": 0.7, "rougeL": 0.7}
    y = 175
    for r in rows:
        model = r["model"]
        draw.text((70, y + 8), model, fill="#1F2A36", font=label_f)
        x0 = 420
        for key, label, color in metrics:
            value = float(r[key])
            normalized = value / max_values[key]
            bar_w = int(680 * min(normalized, 1.0))
            yy = y + {"sacrebleu": 0, "rouge1": 38, "rougeL": 76}[key]
            draw.rounded_rectangle((x0, yy, x0 + 680, yy + 24), radius=12, fill="#EEF3F7")
            draw.rounded_rectangle((x0, yy, x0 + bar_w, yy + 24), radius=12, fill=color)
            text = f"{label}: {value:.4f}" if key != "sacrebleu" else f"{label}: {value:.2f}"
            draw.text((x0 + 710, yy - 3), text, fill="#2D4052", font=small_f)
        y += 124
    img.save(out)
    return out


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Project Report")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = ACCENT

    for line in ["Generative AI and LLMs", "CSE3720"]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(14)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("By").bold = True
    for name, enrollment in TEAM:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{name}\n{enrollment}")
        r.font.size = Pt(12)

    document.add_paragraph()
    for line in [
        "Department of Computer Science and Engineering",
        "School of Engineering and Technology",
        "BML Munjal University",
        "May 2026",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(12)


def add_declaration(document: Document) -> None:
    add_heading(document, "Declaration by the Candidates", 1)
    add_body(
        document,
        f'We hereby declare that the project entitled "{TITLE}" has been carried out to fulfil the partial '
        "requirements for completion of the core-elective course Generative AI and LLMs offered in the 6th "
        "Semester of the Bachelor of Technology program in the Department of Computer Science and Engineering "
        "during AY 2025-26. This work has been carried out by us and submitted to the course instructor "
        "Dr. Soharab Hossain Shaikh. Due acknowledgments have been made to all external tools, libraries, "
        "papers, APIs, and materials used in this project.",
    )
    document.add_paragraph()
    for name, _ in TEAM:
        add_body(document, f"{name} & Signature: __________________________")
    document.add_paragraph()
    add_body(document, "Place: BML Munjal University")
    add_body(document, "Date: ______ May, 2026")


def add_contents(document: Document) -> None:
    add_heading(document, "Contents", 1)
    rows = [
        ["1", "Executive Summary and Introduction", "1"],
        ["2", "Problem Statement and Objectives", "2"],
        ["3", "System Architecture and Methodology", "3-7"],
        ["4", "Dataset, Preprocessing, and Storage", "8-10"],
        ["5", "Models, Baselines, and PEFT Fine-Tuning", "11-13"],
        ["6", "Quantitative Evaluation and Results", "14-16"],
        ["7", "Qualitative Analysis, Guardrails, and UI", "17-19"],
        ["8", "Real-World Applicability and Conclusion", "20"],
        ["9", "References and Appendix", "21+"],
    ]
    add_table(document, ["S. No.", "Section", "Page No."], rows, widths=[0.7, 4.8, 1.1])


def add_intro(document: Document) -> None:
    add_heading(document, "1. Executive Summary and Introduction", 1)
    add_callout(
        document,
        "Project in one line",
        "ScholarSynth AI is a multi-agent GenAI research assistant that retrieves academic papers, indexes them semantically, "
        "generates grounded literature reviews and research-gap analysis, and demonstrates measurable improvement using LoRA fine-tuning.",
        LIGHT_BLUE,
    )
    add_body(
        document,
        "The volume of scientific literature is growing faster than most students and early-stage researchers can manually review. "
        "A user who wants to explore a topic such as retrieval-augmented generation, AI chatbots, guardrails, or transformer models "
        "must search multiple sources, inspect abstracts, compare methods, and then synthesize the material into a coherent review.",
    )
    add_body(
        document,
        "This project addresses that challenge through an end-to-end system for topic-based paper retrieval, semantic search, "
        "RAG-based generation, fine-tuned academic summarization, guardrails, and a Streamlit demonstration interface.",
    )
    nums = project_numbers()
    add_table(
        document,
        ["Key Asset", "Current Status"],
        [
            ["Raw paper dataset", f"{nums['raw_papers']} papers"],
            ["Processed text corpus", f"{nums['processed_chunks']} chunks"],
            ["Fine-tuning dataset", f"{nums['finetune_examples']} instruction examples"],
            ["Vector storage", "ChromaDB persistent index"],
            ["Metadata storage", "SQLite paper database"],
            ["Fine-tuned model", "FLAN-T5-base + LoRA adapter"],
            ["Frontend", "Streamlit app with model selector and guardrail panels"],
        ],
        widths=[2.2, 4.4],
    )


def add_problem_objectives(document: Document) -> None:
    add_heading(document, "2. Problem Statement and Objectives", 1)
    add_body(
        document,
        "Traditional paper search is mostly keyword-driven and does not automatically explain, compare, or synthesize papers. "
        "Generic LLMs can help write text, but they may hallucinate claims, citations, or paper details when not grounded in retrieved evidence.",
    )
    add_heading(document, "2.1 Objectives", 2)
    add_bullets(
        document,
        [
            "Collect an application-specific research paper dataset from arXiv and Semantic Scholar.",
            "Preprocess academic text through cleaning, normalization, deduplication, chunking, and train-validation-test splitting.",
            "Store paper metadata in SQLite and semantic embeddings in ChromaDB.",
            "Build baseline systems using pretrained FLAN-T5, prompt-engineered FLAN-T5, and RAG.",
            "Fine-tune FLAN-T5 with LoRA for academic summarization and explanation tasks.",
            "Compare pretrained, prompt-engineered, RAG, fine-tuned LoRA, and RAG + LoRA systems quantitatively and qualitatively.",
            "Implement input and output guardrails for safer research assistance.",
            "Provide a Streamlit UI suitable for project demonstration and viva explanation.",
        ],
    )
    add_heading(document, "2.2 Evaluation Criteria Mapping", 2)
    add_table(
        document,
        ["Rubric Requirement", "How This Project Satisfies It"],
        [
            ["Dataset quality and split", "5,000 papers, 10,339 chunks, 8,631 fine-tuning examples, 75/10/15 split."],
            ["PEFT fine-tuning", "FLAN-T5-base fine-tuned using LoRA adapter in Google Colab."],
            ["Baseline comparison", "Pretrained, prompt-engineered, RAG, fine-tuned LoRA, and RAG + LoRA compared."],
            ["Data storage", "SQLite stores metadata; ChromaDB stores vector embeddings."],
            ["Quantitative metrics", "BLEU and ROUGE values saved; BERTScore support exists in evaluation code/notebooks."],
            ["Error analysis and guardrails", "Failure summaries, hallucination-risk proxies, input/output guardrail checks."],
            ["Real-world applicability", "Helps students and researchers explore topics, draft reviews, and inspect source evidence."],
        ],
        widths=[2.2, 4.5],
    )


def add_architecture(document: Document) -> None:
    add_heading(document, "3. System Architecture and Methodology", 1)
    add_body(
        document,
        "The system is divided into independent but connected stages. This makes the project easier to explain, test, and extend: "
        "data collection creates the corpus, preprocessing prepares usable chunks, storage enables retrieval, generation produces outputs, "
        "and evaluation demonstrates whether fine-tuning improves over baselines.",
    )
    img = create_workflow_image()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Inches(6.75))
    cap = document.add_paragraph("Figure 1: End-to-end ScholarSynth AI workflow.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    add_table(
        document,
        ["Component", "Role in the System"],
        [
            ["Paper Search", "Fetches titles, abstracts, authors, dates, URLs, and source metadata from arXiv/Semantic Scholar."],
            ["Preprocessing", "Cleans text, removes invalid records, deduplicates papers, chunks content, and prepares splits."],
            ["Vector Store", "Embeds text with all-MiniLM-L6-v2 and persists vectors in ChromaDB."],
            ["Metadata Store", "Keeps structured paper metadata in SQLite for inspection and future filtering."],
            ["RAG Pipeline", "Retrieves relevant chunks and injects them into FLAN-T5 prompts."],
            ["LoRA Adapter", "Adapts FLAN-T5-base to academic tasks without full model fine-tuning."],
            ["Agents", "Separate literature review, gap analysis, explanation, and guardrail responsibilities."],
            ["Streamlit UI", "Presents the system as a usable research-paper explorer and generator."],
        ],
        widths=[1.8, 4.9],
    )


def add_dataset_storage(document: Document) -> None:
    nums = project_numbers()
    add_heading(document, "4. Dataset, Preprocessing, and Storage", 1)
    add_body(
        document,
        "The dataset focuses on Generative AI and LLM-related research areas such as Hugging Face models, AI chatbots, RAG, transformers, "
        "neural networks, hallucination, guardrails, multi-agent systems, scholarly search, and literature review generation.",
    )
    add_heading(document, "4.1 Dataset Size and Split", 2)
    add_table(
        document,
        ["Data Item", "Count"],
        [
            ["Raw papers", nums["raw_papers"]],
            ["Processed chunks", nums["processed_chunks"]],
            ["Fine-tuning examples", nums["finetune_examples"]],
            ["Train chunks", nums["train_chunks"]],
            ["Validation chunks", nums["val_chunks"]],
            ["Test chunks", nums["test_chunks"]],
            ["Fine-tune train/val/test", f"{nums['train_ft']} / {nums['val_ft']} / {nums['test_ft']}"],
        ],
        widths=[2.5, 2.0],
    )
    add_body(
        document,
        "The split ratio is 75% training, 10% validation, and 15% testing. The split is important because fine-tuning and evaluation "
        "must not use the same examples; otherwise, reported improvements would not be reliable.",
    )
    add_heading(document, "4.2 Preprocessing Steps", 2)
    add_numbered(
        document,
        [
            "Remove missing or invalid title/abstract records.",
            "Normalize whitespace and clean academic text.",
            "Deduplicate repeated papers from API overlap.",
            "Combine title and abstract into retrievable text.",
            "Chunk long content for vector retrieval.",
            "Create instruction-style examples for summarization, explanation, Q&A, comparison, literature review, and gap analysis.",
        ],
    )
    add_heading(document, "4.3 Storage Design", 2)
    add_table(
        document,
        ["Storage", "Used For", "Why It Is Needed"],
        [
            ["SQLite", "Paper metadata such as title, authors, abstract, source, URL, and publication date.", "Supports structured access, inspection, and future filtering."],
            ["ChromaDB", "Dense embeddings for processed chunks.", "Supports semantic retrieval beyond keyword matching."],
            ["CSV/JSONL", "Raw dataset, splits, and fine-tuning examples.", "Makes notebooks reproducible and easy to audit."],
        ],
        widths=[1.4, 3.0, 2.3],
    )


def add_models(document: Document) -> None:
    add_heading(document, "5. Models, Baselines, and PEFT Fine-Tuning", 1)
    add_table(
        document,
        ["System", "Model/Method", "Purpose"],
        [
            ["Pretrained baseline", "google/flan-t5-base", "Checks zero/few-shot academic generation quality."],
            ["Prompt-engineered baseline", "google/flan-t5-base with structured task prompt", "Tests prompt design without training."],
            ["RAG system", "ChromaDB retrieval + FLAN-T5-base", "Adds retrieved paper evidence to generation."],
            ["Fine-tuned LoRA", "FLAN-T5-base + LoRA adapter", "Improves academic task formatting and summarization quality."],
            ["RAG + LoRA", "Retrieved context + fine-tuned adapter", "Combines evidence grounding with fine-tuned academic behavior."],
            ["Embedding model", "sentence-transformers/all-MiniLM-L6-v2", "Creates compact semantic embeddings for paper chunks."],
        ],
        widths=[1.8, 2.4, 2.5],
    )
    add_heading(document, "5.1 Why LoRA/PEFT Was Used", 2)
    add_body(
        document,
        "Full fine-tuning of a language model is expensive and impractical on a student MacBook. LoRA is parameter-efficient: it keeps "
        "the base FLAN-T5 model mostly frozen and trains a small adapter. This is suitable for Google Colab, easier to store, and still "
        "demonstrates actual fine-tuning rather than only API prompting.",
    )
    add_heading(document, "5.2 Multi-Agent Responsibilities", 2)
    add_table(
        document,
        ["Agent", "Responsibility"],
        [
            ["LiteratureReviewAgent", "Creates structured reviews from retrieved evidence."],
            ["ResearchGapAgent", "Identifies limitations and possible future directions."],
            ["TechnicalExplainerAgent", "Explains complex research concepts in simpler language."],
            ["GuardrailAgent", "Validates user inputs and generated outputs for safer use."],
        ],
        widths=[2.2, 4.5],
    )


def add_results(document: Document) -> None:
    add_heading(document, "6. Quantitative Evaluation and Results", 1)
    add_body(
        document,
        "The latest saved comparison evaluates five systems on 200 examples. The larger LoRA-only test evaluates the trained adapter on "
        "1,200 examples. BLEU and ROUGE measure overlap with reference outputs; BERTScore support is implemented in the project for semantic scoring.",
    )
    rows = read_csv_dict(ROOT / "outputs/baseline_200_comparison_table.csv")
    rows = sorted(rows, key=lambda r: float(r.get("rougeL", 0)), reverse=True)
    add_table(
        document,
        ["Model", "Examples", "BLEU", "ROUGE-1", "ROUGE-2", "ROUGE-L"],
        [
            [
                r["model"],
                r["eval_examples"],
                f"{float(r['sacrebleu']):.4f}",
                f"{float(r['rouge1']):.4f}",
                f"{float(r['rouge2']):.4f}",
                f"{float(r['rougeL']):.4f}",
            ]
            for r in rows
        ],
        widths=[1.7, 0.8, 0.8, 0.9, 0.9, 0.9],
    )
    chart = create_metrics_chart()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(chart), width=Inches(6.75))
    cap = document.add_paragraph("Figure 2: BLEU, ROUGE-1, and ROUGE-L comparison across evaluated systems.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    lora = read_csv_dict(ROOT / "outputs/lora_test_metrics.csv")
    if lora:
        r = lora[0]
        add_callout(
            document,
            "Best fine-tuned model result",
            f"The LoRA-only test on {r['eval_examples']} examples achieved BLEU {float(r['bleu']):.4f}, "
            f"ROUGE-1 {float(r['rouge1']):.4f}, ROUGE-2 {float(r['rouge2']):.4f}, and ROUGE-L {float(r['rougeL']):.4f}.",
            LIGHT_GREEN,
        )
    add_heading(document, "6.1 Interpretation of Results", 2)
    add_bullets(
        document,
        [
            "Fine-tuned LoRA is the strongest lexical scorer, showing clear improvement over all baseline systems.",
            "RAG + LoRA is second overall and is preferable for demo answers where source grounding matters.",
            "Pretrained and prompt-engineered FLAN-T5 produce weaker academic task outputs because they are not adapted to the project dataset.",
            "RAG alone is useful for evidence retrieval, but generation quality improves substantially when paired with the LoRA adapter.",
        ],
    )
    task_rows = read_csv_dict(ROOT / "outputs/rag_plus_lora_200_task_metrics.csv")
    if task_rows:
        add_heading(document, "6.2 RAG + LoRA Task-Level Results", 2)
        add_table(
            document,
            ["Task", "BLEU", "ROUGE-L"],
            [[r["task"], f"{float(r['bleu']):.4f}", f"{float(r['rougeL']):.4f}"] for r in task_rows],
            widths=[3.0, 1.2, 1.2],
        )


def add_qualitative_guardrails_ui(document: Document) -> None:
    add_heading(document, "7. Qualitative Analysis, Guardrails, and UI", 1)
    add_body(
        document,
        "Quantitative metrics alone cannot determine whether a research assistant is safe or useful. The project therefore includes "
        "qualitative review through best/worst cases, retrieved-evidence inspection, hallucination-risk proxies, and guardrail checks.",
    )
    fail_rows = read_csv_dict(ROOT / "outputs/error_analysis_failure_summary.csv")
    add_table(
        document,
        ["Strategy", "Failure Proxy", "Count"],
        [[r["strategy"], r["failure_type"], r["count"]] for r in fail_rows],
        widths=[2.2, 3.1, 0.9],
    )
    guard_rows = read_csv_dict(ROOT / "outputs/error_analysis_guardrail_summary.csv")
    add_heading(document, "7.1 Guardrail Findings", 2)
    add_table(
        document,
        ["Strategy", "Guardrail Finding", "Count"],
        [[r["strategy"], r["message"], r["count"]] for r in guard_rows],
        widths=[1.7, 4.1, 0.8],
    )
    add_heading(document, "7.2 Input and Output Guardrails", 2)
    add_bullets(
        document,
        [
            "Input guardrails block empty queries, overly long queries, prompt-injection attempts, and credential-like text.",
            "Warnings are shown for very short, off-scope, or personal-data-like queries.",
            "Output guardrails flag empty answers, repetitive wording, weak evidence overlap, generic disclaimers, citation-like text, and unsupported quantitative claims.",
            "The Streamlit UI displays guardrail findings in expandable panels so the demo user can inspect risks rather than blindly trust the answer.",
        ],
    )
    add_heading(document, "7.3 Streamlit Demonstration Interface", 2)
    add_body(
        document,
        "The frontend in app.py provides the demonstration layer. It includes dataset status cards, light/dark theme support, paper fetching, "
        "local index building, task selection, model selection, retrieved evidence tabs, and guardrail feedback. The available generation modes "
        "are Base FLAN-T5, Fine-tuned LoRA, and RAG + Fine-tuned LoRA.",
    )


def add_applicability_conclusion(document: Document) -> None:
    add_heading(document, "8. Real-World Applicability and Conclusion", 1)
    add_body(
        document,
        "ScholarSynth AI is useful for students, research interns, and academic teams who need to quickly understand a new research area. "
        "It can reduce time spent on paper discovery, abstract reading, method comparison, and first-draft literature review preparation. "
        "The system remains explainable because retrieved paper titles and chunks are available for inspection.",
    )
    add_heading(document, "8.1 Viva Explanation Flow", 2)
    add_numbered(
        document,
        [
            "Start with the problem: too many papers, keyword search is limited, and LLMs hallucinate without evidence.",
            "Explain the data pipeline: collect 5,000 papers, clean them, chunk them, and split them properly.",
            "Explain storage: SQLite for metadata and ChromaDB for embeddings.",
            "Explain models: FLAN-T5 baselines, RAG, LoRA fine-tuning, and RAG + LoRA.",
            "Show the result table: LoRA improves strongly over pretrained, prompt-engineered, and RAG-only baselines.",
            "Show guardrails and error analysis: the system detects weak grounding, repetition, and unsupported claims.",
            "End with real-world use: faster literature exploration with visible evidence and a Streamlit demo.",
        ],
    )
    add_heading(document, "8.2 Conclusion", 2)
    add_body(
        document,
        "The project satisfies the main evaluation requirements: application-specific dataset construction, preprocessing and split, PEFT fine-tuning, "
        "baseline comparison, vector and SQL storage, quantitative evaluation, qualitative error analysis, guardrails, and frontend integration. "
        "The clearest improvement is achieved by the fine-tuned LoRA model, while the RAG + LoRA system is the strongest demonstration choice for grounded academic assistance.",
    )


def add_references_appendix(document: Document) -> None:
    add_heading(document, "9. References", 1)
    add_numbered(
        document,
        [
            "Raffel et al., Exploring the Limits of Transfer Learning with a Unified Text-to-Text Transformer, 2020.",
            "Chung et al., Scaling Instruction-Finetuned Language Models, 2022.",
            "Hu et al., LoRA: Low-Rank Adaptation of Large Language Models, 2021.",
            "Lewis et al., Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks, 2020.",
            "Reimers and Gurevych, Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks, 2019.",
            "Zhang et al., BERTScore: Evaluating Text Generation with BERT, 2020.",
            "Papineni et al., BLEU: a Method for Automatic Evaluation of Machine Translation, 2002.",
            "Lin, ROUGE: A Package for Automatic Evaluation of Summaries, 2004.",
            "arXiv API Documentation and Semantic Scholar Graph API Documentation.",
            "Hugging Face Transformers, PEFT, Sentence Transformers, ChromaDB, and Streamlit documentation.",
        ],
    )
    add_heading(document, "Appendix A: Project Files", 1)
    add_table(
        document,
        ["Path", "Purpose"],
        [
            ["notebooks/01_data_collection.ipynb", "Paper collection workflow."],
            ["notebooks/02_vector_database.ipynb", "SQLite and ChromaDB indexing."],
            ["notebooks/04_peft_finetuning_colab.ipynb", "Colab LoRA fine-tuning notebook."],
            ["notebooks/07_finetuned_lora_comparison.ipynb", "Fine-tuned and RAG + LoRA comparison workflow."],
            ["notebooks/08_error_analysis_and_guardrails.ipynb", "Qualitative error analysis and guardrail summaries."],
            ["src/baseline_eval.py", "Evaluation controller for model comparisons."],
            ["src/agents.py", "Multi-agent and guardrail logic."],
            ["app.py", "Streamlit user interface."],
            ["outputs/baseline_200_comparison.md", "Final comparison summary."],
            ["outputs/error_analysis.md", "Failure-case and guardrail report."],
            ["models/flan_t5_lora/", "Fine-tuned LoRA adapter files."],
        ],
        widths=[2.8, 3.9],
    )


def main() -> None:
    document = Document(str(TEMPLATE)) if TEMPLATE.exists() else Document()
    clear_document_body(document)
    setup_document(document)
    add_cover(document)
    document.add_page_break()
    add_declaration(document)
    document.add_page_break()
    add_contents(document)
    document.add_page_break()
    add_intro(document)
    add_problem_objectives(document)
    document.add_page_break()
    add_architecture(document)
    add_dataset_storage(document)
    document.add_page_break()
    add_models(document)
    add_results(document)
    document.add_page_break()
    add_qualitative_guardrails_ui(document)
    add_applicability_conclusion(document)
    add_references_appendix(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
