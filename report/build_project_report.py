from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "report" / "ScholarSynth_AI_Project_Report.docx"
ASSET_DIR = ROOT / "report" / "assets"

TITLE = (
    "Autonomous Research Assistant: A Multi-Agent Generative AI System for "
    "Research Paper Exploration, Literature Review Generation, and Research Gap Analysis"
)

STUDENTS = [
    ("Akshit Wadhwa", "230784"),
    ("Anish Choudhary", "230820"),
    ("Deepnder", "230821"),
]

ACCENT = "1F4E79"
LIGHT = "EAF2FB"
DARK = RGBColor(31, 78, 121)
TABLE_COUNTER = 0


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = str(text).split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_rounded_rect(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], fill: str, outline: str | None = None) -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=2 if outline else 1)


def create_flowchart_image() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "project_flowchart.png"
    width, height = 1800, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(46, bold=True)
    subtitle_font = load_font(24)
    box_font = load_font(22, bold=True)
    small_font = load_font(18)

    draw.rectangle((0, 0, width, height), fill="#F7FBFF")
    draw.text((70, 42), "ScholarSynth AI: Project Workflow", fill="#173B63", font=title_font)
    draw.text(
        (72, 100),
        "Research paper retrieval, semantic storage, RAG generation, LoRA fine-tuning, evaluation, and UI",
        fill="#496A88",
        font=subtitle_font,
    )

    boxes = [
        (80, 200, 350, 330, "#DFF3FF", "User Topic", "Research query from Streamlit"),
        (430, 200, 740, 330, "#E8E4FF", "Paper APIs", "arXiv + Semantic Scholar"),
        (820, 200, 1160, 330, "#E8F8EC", "Dataset", "5,000 papers, 10,339 chunks"),
        (1240, 200, 1630, 330, "#FFF1D6", "Preprocessing", "Clean, deduplicate, chunk, split"),
        (200, 470, 520, 600, "#E8F8EC", "Storage", "SQLite metadata + ChromaDB vectors"),
        (600, 470, 920, 600, "#DFF3FF", "Retrieval", "all-MiniLM-L6-v2 embeddings"),
        (1000, 470, 1320, 600, "#E8E4FF", "Generation", "FLAN-T5-base + RAG context"),
        (1400, 470, 1700, 600, "#FFE5EA", "LoRA", "PEFT fine-tuned adapter"),
        (440, 720, 760, 830, "#F3EFFF", "Multi-Agent Layer", "Review, gap, explainer, guardrail"),
        (850, 720, 1170, 830, "#E6F7F2", "Outputs", "Review, Q&A, gaps, explanation"),
        (1260, 720, 1600, 830, "#FFF1D6", "Evaluation", "BLEU, ROUGE, BERTScore, errors"),
    ]

    for x1, y1, x2, y2, color, heading, body in boxes:
        draw_rounded_rect(draw, (x1, y1, x2, y2), color, "#B8C7D6")
        draw.text((x1 + 22, y1 + 22), heading, fill="#173B63", font=box_font)
        for i, line in enumerate(wrap_text(draw, body, small_font, x2 - x1 - 44)[:3]):
            draw.text((x1 + 22, y1 + 65 + i * 24), line, fill="#35536E", font=small_font)

    arrows = [
        ((355, 265), (425, 265)),
        ((745, 265), (815, 265)),
        ((1165, 265), (1235, 265)),
        ((1435, 335), (1445, 465)),
        ((1240, 535), (925, 535)),
        ((595, 535), (525, 535)),
        ((925, 535), (995, 535)),
        ((1325, 535), (1395, 535)),
        ((1548, 605), (1120, 715)),
        ((850, 665), (760, 725)),
        ((765, 775), (845, 775)),
        ((1175, 775), (1255, 775)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#2B6EA6", width=5)
        ex, ey = end
        sx, sy = start
        direction = 1 if ex >= sx else -1
        draw.polygon([(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)], fill="#2B6EA6")

    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    global TABLE_COUNTER
    TABLE_COUNTER += 1
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    output_path = ASSET_DIR / f"table_{TABLE_COUNTER:02d}.png"

    image_width = 1500
    margin_x = 42
    padding_x = 22
    padding_y = 14
    header_font = load_font(27, bold=True)
    body_font = load_font(25)
    caption_font = load_font(20)
    line_gap = 7

    if widths:
        total = sum(widths)
        col_widths = [int((image_width - 2 * margin_x) * width / total) for width in widths]
    else:
        col_widths = [int((image_width - 2 * margin_x) / len(headers))] * len(headers)
    col_widths[-1] += image_width - 2 * margin_x - sum(col_widths)

    scratch = Image.new("RGB", (image_width, 200), "white")
    scratch_draw = ImageDraw.Draw(scratch)

    def row_height(values: list[str], font: ImageFont.FreeTypeFont) -> int:
        max_lines = 1
        for value, col_width in zip(values, col_widths):
            lines = wrap_text(scratch_draw, value, font, col_width - 2 * padding_x)
            max_lines = max(max_lines, len(lines))
        line_height = font.getbbox("Ag")[3] - font.getbbox("Ag")[1] + line_gap
        return max(58, padding_y * 2 + max_lines * line_height)

    heights = [row_height(headers, header_font)] + [row_height(row, body_font) for row in rows]
    image_height = sum(heights) + 2 * margin_x
    img = Image.new("RGB", (image_width, image_height), "white")
    draw = ImageDraw.Draw(img)

    y = margin_x
    all_rows = [headers] + rows
    for row_index, (values, height) in enumerate(zip(all_rows, heights)):
        x = margin_x
        is_header = row_index == 0
        fill = "#1F4E79" if is_header else ("#F7FBFF" if row_index % 2 == 0 else "#FFFFFF")
        font = header_font if is_header else body_font
        text_fill = "#FFFFFF" if is_header else "#1F2933"
        for value, col_width in zip(values, col_widths):
            draw.rectangle((x, y, x + col_width, y + height), fill=fill, outline="#B8C7D6", width=2)
            lines = wrap_text(draw, value, font, col_width - 2 * padding_x)
            line_height = font.getbbox("Ag")[3] - font.getbbox("Ag")[1] + line_gap
            text_y = y + (height - len(lines) * line_height) // 2
            for line in lines:
                draw.text((x + padding_x, text_y), line, fill=text_fill, font=font)
                text_y += line_height
            x += col_width
        y += height

    draw.rectangle((margin_x, margin_x, image_width - margin_x, image_height - margin_x), outline="#8AA6BF", width=3)
    img.save(output_path)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(output_path), width=Inches(6.75))
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    return output_path


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = DARK


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def setup_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = DARK
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def add_cover(document: Document) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(8)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK

    for line in ["Generative AI and LLMs", "CSE3720"]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(14)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("By").bold = True

    for name, enrollment in STUDENTS:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{name}\n{enrollment}")
        run.font.size = Pt(12)

    document.add_paragraph()
    for line in [
        "Department of Computer Science and Engineering",
        "School of Engineering and Technology",
        "BML Munjal University",
        "May 2026",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)


def add_declaration(document: Document) -> None:
    add_heading(document, "Declaration by the Candidates", 1)
    add_body(
        document,
        f'We hereby declare that the project entitled "{TITLE}" has been carried out to fulfil the partial '
        "requirements for completion of the core-elective course Generative AI and LLMs offered in the 6th "
        "Semester of the Bachelor of Technology (B.Tech) program in the Department of Computer Science and "
        "Engineering during AY-2025-26 (even semester). This experimental work has been carried out by us and "
        "submitted to the course instructor Dr. Soharab Hossain Shaikh. Due acknowledgments have been made in "
        "the text of the project to all other materials used. This project has been prepared in full compliance "
        "with the requirements and constraints of the prescribed curriculum.",
    )
    document.add_paragraph()
    for name, _ in STUDENTS:
        add_body(document, f"{name} & Signature: __________________________")
    document.add_paragraph()
    add_body(document, "Place: BML Munjal University")
    add_body(document, "Date: ______ May, 2026")


def add_contents(document: Document) -> None:
    add_heading(document, "Contents", 1)
    rows = [
        ["1", "Introduction", "1-2"],
        ["2", "Problem Statement", "3-4"],
        ["3", "Literature Review", "5-10"],
        ["4", "Methodology", "11-15"],
        ["5", "Technology Stack", "16"],
        ["6", "Results", "17-23"],
        ["7", "Conclusions", "24"],
        ["8", "References", ""],
        ["9", "Appendix", ""],
    ]
    add_table(document, ["S. No.", "Section", "Page No."], rows, widths=[0.7, 4.8, 1.2])


def add_introduction(document: Document) -> None:
    add_heading(document, "1. Introduction", 1)
    add_body(
        document,
        "The volume of scientific literature is increasing rapidly, making it difficult for students and "
        "researchers to discover relevant papers, understand technical methods, and identify open research "
        "problems. Traditional keyword-based search often misses semantically related papers and does not "
        "directly help users synthesize information into a structured literature review.",
    )
    add_body(
        document,
        "This project, ScholarSynth AI, is an autonomous research assistant that combines research paper "
        "retrieval, semantic search, Retrieval-Augmented Generation (RAG), PEFT-based LoRA fine-tuning, "
        "multi-agent task separation, quantitative evaluation, and a Streamlit-based user interface. The "
        "system supports topic-based paper exploration, literature review generation, paper question answering, "
        "research gap detection, and simplified technical explanation.",
    )
    add_body(
        document,
        "The solution is designed for academic use cases where users need grounded answers based on retrieved "
        "research evidence rather than unsupported or hallucinated model outputs. It uses arXiv and Semantic "
        "Scholar for paper collection, ChromaDB for vector search, SQLite for metadata storage, FLAN-T5 for "
        "generation, and LoRA for lightweight domain adaptation.",
    )


def add_problem_statement(document: Document) -> None:
    add_heading(document, "2. Problem Statement", 1)
    add_body(
        document,
        "Researchers often spend significant time searching for relevant academic papers, reading abstracts, "
        "comparing methods, and writing literature reviews. This work becomes harder when the topic is new, "
        "interdisciplinary, or contains rapidly changing terminology such as large language models, RAG, "
        "agentic systems, and AI guardrails.",
    )
    add_body(
        document,
        "The project addresses the following problems:",
    )
    add_bullets(
        document,
        [
            "Keyword search does not always retrieve papers that are semantically related to the user topic.",
            "Manual reading and comparison of many papers is time-consuming.",
            "Large language models can hallucinate paper titles, claims, datasets, and citations.",
            "Generic pretrained models may not produce strong academic summaries or research gap analysis.",
            "Students need a simpler interface for exploring papers and understanding technical methods.",
        ],
    )
    add_body(
        document,
        "The goal is to build a grounded AI research assistant that retrieves relevant papers, stores them "
        "systematically, generates evidence-aware outputs, and demonstrates improvement using proper baseline "
        "comparison and LoRA fine-tuning.",
    )


def add_literature_review(document: Document) -> None:
    add_heading(document, "3. Literature Review", 1)
    sections = [
        (
            "3.1 Large Language Models",
            "Large language models have shown strong performance in summarization, question answering, and "
            "text generation. However, pretrained models may generate generic or unsupported content when they "
            "are not grounded in external evidence. For academic tasks, this creates a risk of hallucinated "
            "citations and inaccurate claims.",
        ),
        (
            "3.2 Retrieval-Augmented Generation",
            "RAG combines information retrieval with generation. Instead of asking the model to answer only "
            "from its internal parameters, relevant documents are retrieved and inserted into the prompt. This "
            "improves factual grounding and makes the answer easier to verify.",
        ),
        (
            "3.3 Semantic Search and Vector Databases",
            "Semantic search represents text as dense embeddings and retrieves documents based on meaning. "
            "Vector databases such as ChromaDB support efficient similarity search over these embeddings. "
            "This is useful for research paper exploration because related papers may use different keywords.",
        ),
        (
            "3.4 PEFT and LoRA",
            "Parameter-Efficient Fine-Tuning (PEFT) updates a small number of trainable parameters instead of "
            "fine-tuning the complete model. LoRA injects low-rank adapter matrices into transformer layers, "
            "making it practical to fine-tune models on limited GPU resources such as Google Colab.",
        ),
        (
            "3.5 Evaluation of Generated Text",
            "Automatic metrics such as BLEU, ROUGE, and BERTScore are commonly used to compare generated text "
            "against reference outputs. BLEU measures n-gram overlap, ROUGE is widely used for summarization, "
            "and BERTScore compares semantic similarity using contextual embeddings.",
        ),
        (
            "3.6 Guardrails for AI Systems",
            "Guardrails are validation checks around LLM systems. In a research assistant, guardrails should "
            "block invalid inputs, reduce unsupported claims, prevent fake citations, and instruct the system "
            "to state when evidence is limited.",
        ),
    ]
    for heading, text in sections:
        add_heading(document, heading, 2)
        add_body(document, text)


def add_methodology(document: Document) -> None:
    add_heading(document, "4. Methodology", 1)
    add_heading(document, "4.1 System Architecture", 2)
    add_body(
        document,
        "The system follows a modular pipeline. Papers are collected from external APIs, cleaned and chunked, "
        "embedded using a sentence-transformer model, stored in ChromaDB and SQLite, retrieved using semantic "
        "search, and passed to FLAN-T5 for generation. Multi-agent components separate literature review, "
        "research gap, technical explanation, and guardrail responsibilities.",
    )
    flowchart_path = create_flowchart_image()
    figure = document.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.add_run().add_picture(str(flowchart_path), width=Inches(6.75))
    caption = document.add_paragraph("Figure 1: End-to-end ScholarSynth AI workflow.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    add_table(
        document,
        ["Stage", "Description"],
        [
            ["Data Collection", "Retrieve paper metadata and abstracts from arXiv and Semantic Scholar."],
            ["Preprocessing", "Clean text, remove duplicates, normalize fields, and chunk title plus abstract text."],
            ["Storage", "Store metadata in SQLite and vector embeddings in ChromaDB."],
            ["Retrieval", "Use semantic search to retrieve relevant paper chunks for a user query."],
            ["Generation", "Use FLAN-T5 and RAG prompts to generate summaries, reviews, answers, and gaps."],
            ["Fine-Tuning", "Apply LoRA to FLAN-T5 using academic task examples."],
            ["Evaluation", "Compare systems using BLEU, ROUGE, and BERTScore."],
        ],
        widths=[1.8, 4.9],
    )
    add_heading(document, "4.2 Dataset Collection", 2)
    add_body(
        document,
        "The dataset contains papers from topics related to LLMs, Hugging Face models, AI chatbots, RAG, "
        "transformers, neural networks, scholarly search, guardrails, and text generation evaluation. The "
        "collection process used arXiv and Semantic Scholar APIs and saved paper metadata in CSV format.",
    )
    add_table(
        document,
        ["Dataset File", "Count"],
        [
            ["Raw papers", "5,000 papers"],
            ["Processed chunks", "10,339 chunks"],
            ["Fine-tuning examples", "8,631 examples"],
            ["Topics covered", "94 topics"],
            ["Source split", "4,673 arXiv papers and 327 Semantic Scholar papers"],
        ],
        widths=[2.4, 4.3],
    )
    add_heading(document, "4.3 Preprocessing and Data Split", 2)
    add_body(
        document,
        "Preprocessing removes invalid abstracts, deduplicates papers by title and abstract, normalizes "
        "whitespace, combines title and abstract, and chunks long academic text into smaller passages for "
        "retrieval. The final train-validation-test split is 75 percent, 10 percent, and 15 percent.",
    )
    add_table(
        document,
        ["Split", "Chunk Rows", "Fine-Tuning Examples"],
        [
            ["Training", "7,754", "6,473"],
            ["Validation", "1,034", "863"],
            ["Testing", "1,551", "1,295"],
        ],
        widths=[1.7, 2.2, 2.2],
    )
    add_heading(document, "4.4 Fine-Tuning Dataset Tasks", 2)
    add_bullets(
        document,
        [
            "Paper summarization from title and abstract.",
            "Technical explanation for early-stage researchers.",
            "Evidence-based question answering using only the given abstract.",
            "Short literature review generation from grouped papers.",
            "Research gap analysis from retrieved papers.",
            "Comparative analysis across related papers.",
        ],
    )
    add_heading(document, "4.5 Multi-Agent Design", 2)
    add_table(
        document,
        ["Agent", "Responsibility"],
        [
            ["LiteratureReviewAgent", "Generates structured literature reviews from retrieved evidence."],
            ["ResearchGapAgent", "Identifies limitations, gaps, and future research directions."],
            ["TechnicalExplainerAgent", "Explains complex methods and models in simpler language."],
            ["GuardrailAgent", "Validates input and output to reduce invalid queries and unsupported responses."],
        ],
        widths=[2.2, 4.5],
    )


def add_tech_stack(document: Document) -> None:
    add_heading(document, "5. Technology Stack", 1)
    add_table(
        document,
        ["Component", "Technology Used", "Purpose"],
        [
            ["Programming Language", "Python", "Core implementation and notebooks."],
            ["Frontend", "Streamlit", "Interactive user interface."],
            ["Paper Sources", "arXiv API, Semantic Scholar API", "Research paper retrieval."],
            ["Metadata DB", "SQLite", "Structured paper metadata storage."],
            ["Vector DB", "ChromaDB", "Persistent semantic search over embeddings."],
            ["Embedding Model", "sentence-transformers/all-MiniLM-L6-v2", "Generate dense text embeddings."],
            ["Base Model", "google/flan-t5-base", "Text generation baseline."],
            ["Fine-Tuning", "PEFT LoRA", "Parameter-efficient domain adaptation."],
            ["Training Platform", "Google Colab", "GPU-based fine-tuning."],
            ["Evaluation", "BLEU, ROUGE, BERTScore", "Quantitative generation comparison."],
        ],
        widths=[1.7, 2.4, 2.6],
    )


def add_results(document: Document) -> None:
    add_heading(document, "6. Results", 1)
    add_heading(document, "6.1 Dataset and Storage Results", 2)
    add_body(
        document,
        "The project successfully created a large application-specific academic dataset and rebuilt the "
        "retrieval stores. SQLite contains 5,000 metadata rows, and ChromaDB contains 10,339 indexed chunks. "
        "This satisfies the requirement for both structured data storage and vector database storage.",
    )
    add_heading(document, "6.2 Baseline Evaluation", 2)
    add_body(
        document,
        "A 200-example baseline evaluation was completed for pretrained FLAN-T5, prompt-engineered FLAN-T5, "
        "and the RAG system. The metrics are shown below.",
    )
    add_table(
        document,
        ["Model", "BLEU", "ROUGE-1", "ROUGE-2", "ROUGE-L", "BERTScore F1"],
        [
            ["pretrained", "0.0117", "0.2100", "0.1011", "0.1675", "0.7849"],
            ["rag_system", "0.0116", "0.1649", "0.0749", "0.1355", "0.7632"],
            ["prompt_engineered", "0.0026", "0.1384", "0.0727", "0.1200", "0.7537"],
        ],
        widths=[1.6, 0.9, 1.0, 1.0, 1.0, 1.2],
    )
    add_body(
        document,
        "The pretrained model scores highest on lexical overlap because many reference answers are derived "
        "from paper abstracts. The RAG system remains important because it provides retrieved evidence and "
        "supports grounded source-aware generation, even when lexical overlap metrics do not fully capture "
        "that benefit.",
    )
    add_heading(document, "6.3 Fine-Tuned LoRA Model", 2)
    add_body(
        document,
        "The PEFT fine-tuned LoRA adapter has been trained in Google Colab and saved locally under "
        "models/flan_t5_lora. The adapter uses FLAN-T5-base as the base model with LoRA target modules q and v. "
        "A final all-system evaluation notebook has been prepared to compare pretrained, prompt-engineered, "
        "RAG, fine_tuned_lora, and rag_plus_lora systems. The final table should be updated after running the "
        "200-example LoRA comparison notebook.",
    )
    add_table(
        document,
        ["System", "Status"],
        [
            ["pretrained FLAN-T5", "Baseline evaluated on 200 examples."],
            ["prompt-engineered FLAN-T5", "Baseline evaluated on 200 examples."],
            ["RAG system", "Baseline evaluated on 200 examples."],
            ["fine-tuned LoRA model", "Adapter trained; final 200-example evaluation pending."],
            ["RAG + LoRA model", "Notebook prepared; final 200-example evaluation pending."],
        ],
        widths=[2.4, 4.3],
    )
    add_heading(document, "6.4 Qualitative and Error Analysis", 2)
    add_body(
        document,
        "Qualitative analysis focuses on whether outputs are grounded, specific, non-repetitive, and useful "
        "for a research user. The main observed and expected failure modes are listed below.",
    )
    add_table(
        document,
        ["Failure Case", "Description", "Mitigation"],
        [
            ["Hallucinated citation", "Model may invent paper titles or unsupported claims.", "Use RAG evidence and fake-citation checks."],
            ["Irrelevant retrieval", "Vector search may retrieve loosely related papers.", "Use topic-aware query construction and top-k review."],
            ["Generic answer", "Pretrained model may produce broad academic text.", "Use LoRA fine-tuning and task-specific prompts."],
            ["Weak gap analysis", "Model may list obvious gaps without evidence.", "Require retrieved evidence and source titles."],
            ["Repetition", "Generation may repeat phrases on long prompts.", "Clean generations and limit answer length."],
        ],
        widths=[1.6, 2.6, 2.5],
    )
    add_heading(document, "6.5 Guardrails", 2)
    add_body(
        document,
        "The GuardrailAgent validates user queries and generated outputs. Current checks reject empty queries, "
        "limit overly long inputs, and provide a layer for blocking unsupported outputs. Additional guardrails "
        "for final deployment include preventing fake citations, requiring retrieved evidence for paper claims, "
        "and stating that evidence is limited when retrieval quality is weak.",
    )
    add_heading(document, "6.6 Frontend Integration", 2)
    add_body(
        document,
        "A Streamlit interface is included through app.py. The interface is intended to allow users to enter "
        "a research topic, retrieve relevant papers, generate literature reviews, ask paper-related questions, "
        "detect research gaps, and inspect retrieved sources.",
    )


def add_conclusion(document: Document) -> None:
    add_heading(document, "7. Conclusions", 1)
    add_body(
        document,
        "The project demonstrates an end-to-end GenAI research assistant that goes beyond simple prompting. "
        "It includes application-specific data collection, preprocessing, structured and vector storage, "
        "semantic retrieval, RAG generation, multi-agent task separation, PEFT fine-tuning, quantitative "
        "evaluation, qualitative error analysis, guardrails, and a frontend demonstration path.",
    )
    add_body(
        document,
        "The system is applicable to students, researchers, and academic teams who need to explore a new "
        "research area quickly. It can reduce manual effort in paper discovery, abstract reading, literature "
        "review drafting, and research gap identification. Future work should improve retrieval reranking, "
        "add stricter citation verification, complete the final LoRA comparison, and expand UI support for "
        "saving reports and exporting literature reviews.",
    )


def add_references(document: Document) -> None:
    add_heading(document, "8. References", 1)
    refs = [
        "Raffel et al., Exploring the Limits of Transfer Learning with a Unified Text-to-Text Transformer, 2020.",
        "Chung et al., Scaling Instruction-Finetuned Language Models, 2022.",
        "Hu et al., LoRA: Low-Rank Adaptation of Large Language Models, 2021.",
        "Lewis et al., Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks, 2020.",
        "Reimers and Gurevych, Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks, 2019.",
        "Zhang et al., BERTScore: Evaluating Text Generation with BERT, 2020.",
        "Papineni et al., BLEU: a Method for Automatic Evaluation of Machine Translation, 2002.",
        "Lin, ROUGE: A Package for Automatic Evaluation of Summaries, 2004.",
        "arXiv API Documentation.",
        "Semantic Scholar Graph API Documentation.",
        "ChromaDB Documentation.",
        "Hugging Face Transformers and PEFT Documentation.",
    ]
    add_numbered(document, refs)


def add_appendix(document: Document) -> None:
    add_heading(document, "9. Appendix", 1)
    add_heading(document, "Appendix A: Project Structure", 2)
    add_body(
        document,
        "The repository is organized into data, notebooks, source modules, outputs, models, and report folders. "
        "The notebooks demonstrate the workflow, while the source files contain reusable implementation logic.",
    )
    add_table(
        document,
        ["File or Folder", "Purpose"],
        [
            ["data/", "Raw papers, processed chunks, train/validation/test splits, and fine-tuning JSONL files."],
            ["notebooks/", "Step-wise notebooks for data collection, vector DB, baseline comparison, fine-tuning, and final comparison."],
            ["src/paper_search.py", "arXiv and Semantic Scholar paper retrieval."],
            ["src/preprocessing.py", "Cleaning, chunking, splitting, and fine-tuning data creation."],
            ["src/vector_store.py", "SQLite metadata storage and ChromaDB vector indexing."],
            ["src/rag_pipeline.py", "RAG prompt construction and FLAN-T5 generation."],
            ["src/agents.py", "Multi-agent components and guardrails."],
            ["src/evaluation.py", "BLEU, ROUGE, and BERTScore metric functions."],
            ["src/baseline_eval.py", "Final model comparison controller."],
            ["app.py", "Streamlit frontend."],
        ],
        widths=[2.3, 4.4],
    )
    add_heading(document, "Appendix B: Output Artifacts", 2)
    add_bullets(
        document,
        [
            "outputs/baseline_200_eval_data.csv",
            "outputs/baseline_200_metrics.csv",
            "outputs/baseline_200_generations.csv",
            "outputs/baseline_200_comparison_table.csv",
            "outputs/baseline_200_comparison.md",
            "models/flan_t5_lora/",
            "notebooks/07_finetuned_lora_comparison.ipynb",
        ],
    )


def main() -> None:
    document = Document()
    setup_styles(document)
    add_cover(document)
    add_page_break(document)
    add_declaration(document)
    add_page_break(document)
    add_contents(document)
    add_page_break(document)
    add_introduction(document)
    add_problem_statement(document)
    add_literature_review(document)
    add_methodology(document)
    add_tech_stack(document)
    add_results(document)
    add_conclusion(document)
    add_references(document)
    add_appendix(document)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
